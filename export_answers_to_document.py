from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os


def add_question_and_answer_to_document(path, question, answer):
    # Add an empty line if document already contains content
    if os.path.exists(path):
        # Open existing document
        existing_doc = Document(path)

        # Go to the last paragraph
        last_paragraph = existing_doc.paragraphs[-1]

        # Add an empty line if document already contains content
        if len(existing_doc.paragraphs) > 0:
            existing_doc.add_paragraph()
    else:
        existing_doc = Document()

    question_paragraph = existing_doc.add_paragraph()
    question_run = question_paragraph.add_run(question)
    question_run.bold = True
    question_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    question_paragraph.style = "Heading 2"

    # Set font to Verdana
    question_font = question_run.font
    question_font.name = 'Verdana'
    question_font.size = Pt(12)

    # Add answer as normal text
    answer_paragraph = existing_doc.add_paragraph(answer)
    answer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Set font to Verdana
    answer_runs = answer_paragraph.runs
    for run in answer_runs:
        run.font.name = 'Verdana'
        run.font.size = Pt(12)

    existing_doc.save(path)